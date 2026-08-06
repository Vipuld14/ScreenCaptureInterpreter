"""Analysis engine — multi-image explain + .docx export.

Shared brain used by hotkey_capture.py. Sends all images in ONE API call so
Claude analyses them as a single continuous document, returning a JSON object
with an "explanation" (overview) and "extracted_text" (structure-preserving
Markdown). build_docx() turns the extracted text into a Word document.

Public API:
  load_env()                  -> load ANTHROPIC_API_KEY from .env
  analyse_images(client, ps)  -> {"explanation": str, "extracted_text": str}
  build_docx(result)          -> python-docx Document
"""

import base64
import hashlib
import itertools
import json
import sys
import threading
import time
from pathlib import Path

MODEL = "claude-sonnet-4-6"
# Per-image extraction is an OCR-like task — use a cheaper/faster model to cut cost.
# Reasoning steps (classify, fix) keep MODEL. Change if this model isn't available.
EXTRACT_MODEL = "claude-haiku-4-5-20251001"

# Strict JSON response keeps explanation and extracted text cleanly separated.
SYSTEM_PROMPT = (
    "You are a screen-reading assistant that analyses a sequence of screenshots as one continuous piece of content.\n"
    "The images are ordered and together represent a single document, page, or screen flow.\n"
    "Return a JSON object with exactly two keys:\n"
    "\n"
    '  "explanation": A detailed plain-English overview of what is shown across all the images combined. '
    "Lead with the content type (e.g. 'Spreadsheet:', 'Document:', 'Code editor:'), then describe "
    "the full picture — layout, key elements, how the images relate to each other.\n"
    "\n"
    '  "extracted_text": All visible text from all images stitched together in order as one '
    "continuous document, preserving the original structure throughout. Use Markdown:\n"
    "    - Headings → # / ## / ### etc.\n"
    "    - Bullet lists → -\n"
    "    - Numbered lists → 1. 2. 3.\n"
    "    - Plain paragraphs → plain paragraphs separated by blank lines\n"
    "    Continue structure naturally across images — do not restart or add separators.\n"
    '    If there is no meaningful text, use an empty string "".\n'
    "\n"
    "Return ONLY the raw JSON object. No code fences, no extra keys, no commentary."
)

USER_PROMPT = "Analyse all these screenshots as one continuous document and return the JSON as instructed."


def load_env() -> None:
    """Load ANTHROPIC_API_KEY. Searches, in order: the existing environment, the
    working directory's .env (dev), a .env next to the executable/bundle, and
    ~/.code_capture/.env (used by the installed .app). First hit wins."""
    import os
    if os.environ.get("ANTHROPIC_API_KEY"):
        return
    try:
        from dotenv import load_dotenv
    except ImportError:
        return
    import sys
    from pathlib import Path
    load_dotenv()  # 1) cwd / project .env (development)
    if os.environ.get("ANTHROPIC_API_KEY"):
        return
    for cand in (Path(sys.executable).resolve().parent / ".env",   # 2) next to the app binary
                 Path.home() / ".code_capture" / ".env"):          # 3) installed-app config
        try:
            if cand.exists():
                load_dotenv(cand)
                if os.environ.get("ANTHROPIC_API_KEY"):
                    return
        except Exception:  # noqa: BLE001
            pass


# ── Spinner ────────────────────────────────────────────────────────────────────

PROCESSING_MESSAGES = [
    "Reading pixel data across all images...",
    "Identifying content types and layout structures...",
    "Cross-referencing text regions between images...",
    "Analysing visual hierarchy and document flow...",
    "Stitching content together into a single document...",
    "Extracting and preserving text formatting...",
    "Resolving structure across image boundaries...",
    "Almost there — finalising the analysis...",
]


def _spinner(stop_event: threading.Event) -> None:
    spinner = itertools.cycle(["⠋", "⠙", "⠹", "⠸", "⠼", "⠴", "⠦", "⠧", "⠇", "⠏"])
    messages = itertools.cycle(PROCESSING_MESSAGES)
    current_msg = next(messages)
    msg_timer = time.time()

    while not stop_event.is_set():
        print(f"\r  {next(spinner)}  {current_msg}   ", end="", flush=True)
        time.sleep(0.1)
        if time.time() - msg_timer > 3:
            current_msg = next(messages)
            msg_timer = time.time()

    print("\r" + " " * 70 + "\r", end="", flush=True)


# ── API call ───────────────────────────────────────────────────────────────────

# All images go in one API call so Claude has full cross-image context.
# Streaming is hidden — the spinner keeps the user informed instead.
def analyse_images(client, image_paths: list) -> dict:
    content = []
    for path in image_paths:
        b64 = base64.standard_b64encode(path.read_bytes()).decode()
        content.append({
            "type": "image",
            "source": {"type": "base64", "media_type": _media_type(path), "data": b64},
        })
    content.append({"type": "text", "text": USER_PROMPT})

    stop_event = threading.Event()
    spinner_thread = threading.Thread(target=_spinner, args=(stop_event,), daemon=True)
    spinner_thread.start()

    raw = ""
    try:
        with client.messages.stream(
            model=MODEL,
            max_tokens=4096,
            system=SYSTEM_PROMPT,
            messages=[{"role": "user", "content": content}],
        ) as stream:
            for text in stream.text_stream:
                raw += text
    finally:
        stop_event.set()
        spinner_thread.join()

    raw = raw.strip()

    # Attempt 1: parse directly.
    try:
        return json.loads(raw)
    except json.JSONDecodeError:
        pass

    # Attempt 2: find the outermost { } block in case Claude added preamble text.
    start, end = raw.find("{"), raw.rfind("}")
    if start != -1 and end > start:
        try:
            return json.loads(raw[start:end + 1])
        except json.JSONDecodeError:
            pass

    print("\n[Warning] Could not parse structured response — showing raw output.", file=sys.stderr)
    return {"explanation": raw, "extracted_text": ""}


def _media_type(path: Path) -> str:
    return {
        ".jpg": "image/jpeg", ".jpeg": "image/jpeg",
        ".png": "image/png", ".gif": "image/gif", ".webp": "image/webp",
    }.get(path.suffix.lower(), "image/png")



# ── Incremental (per-image) pipeline ─────────────────────────────────────────────
#
# For long sessions, sending every image in one call is slow, costly, and can
# overflow the response (truncated JSON). Instead we read ONE image at a time,
# cache each result by content hash (so identical frames are never re-read),
# stitch the per-image text together locally, then make a single cheap text-only
# call for the overview. This is the Milestone 5 path used by the hotkey tool.

EXTRACT_SYSTEM_PROMPT = (
    "You are a screen-reading assistant. Transcribe ALL visible text from this single "
    "screenshot EXACTLY as it appears — character for character, space for space, line for line. "
    "Output PLAIN TEXT only. Do NOT convert anything to Markdown and do NOT add any Markdown that "
    "is not literally on screen: no # headings, no ``` code fences, no - bullets, no ** bold. "
    "Just reproduce the raw text/code as-is. "
    "Do NOT include the editor's line-number gutter, code-folding arrows or chevrons, breakpoint "
    "dots, diff/git markers, minimaps, scrollbars, tab bars, or status bars — transcribe only the "
    "actual content. Omit line numbers entirely but keep the code's own indentation. "
    "If several windows or apps are visible, transcribe ONLY the primary code or document content "
    "(the focused editor pane); ignore other windows, browser tabs, the dock, and menu bars. "
    "Reproduce mistakes faithfully: if a line is wrongly indented, mis-aligned, or has a typo, "
    "reproduce it AS-IS. Do NOT correct, complete, reformat, or improve anything. "
    "Example: if a method's def line is at the SAME indentation as its class header (not indented "
    "under it), copy it at that same wrong indentation — do not fix it. "
    "If a line is cut off at the screen edge or unreadable, transcribe what is visible and append "
    "the marker [CUT OFF] — never guess. If there is no meaningful text, output nothing."
)

FINALIZE_SYSTEM_PROMPT = (
    "You are given the full text extracted from a sequence of screenshots that "
    "together form one document. Classify it and summarise it.\n"
    "Return ONLY a JSON object with exactly these keys:\n"
    '  "overview": a concise plain-English overview, leading with the content type.\n'
    '  "is_code": true if the content is primarily source code, otherwise false.\n'
    '  "language": if code, the programming language name (e.g. "Python", "C++", '
    '"C#", "JavaScript"); otherwise "".\n'
    '  "extension": if code, the conventional source-file extension WITHOUT a dot '
    '(e.g. "py", "cpp", "cs", "js"); otherwise "".\n'
    "Return only the raw JSON object — no code fences, no commentary."
)


def extract_one(client, path: Path) -> str:
    """Send ONE image to the API and return its visible text as Markdown."""
    b64 = base64.standard_b64encode(path.read_bytes()).decode()
    msg = client.messages.create(
        model=EXTRACT_MODEL,
        max_tokens=4096,
        system=EXTRACT_SYSTEM_PROMPT,
        messages=[{"role": "user", "content": [
            {"type": "image", "source": {"type": "base64", "media_type": _media_type(path), "data": b64}},
            {"type": "text", "text": "Extract the text from this screenshot as Markdown."},
        ]}],
    )
    return "".join(getattr(b, "text", "") for b in msg.content).strip()


def _parse_json(raw: str) -> dict:
    """Best-effort JSON parse: direct, then the outermost { } block."""
    try:
        return json.loads(raw)
    except json.JSONDecodeError:
        pass
    start, end = raw.find("{"), raw.rfind("}")
    if start != -1 and end > start:
        try:
            return json.loads(raw[start:end + 1])
        except json.JSONDecodeError:
            pass
    return None


def synthesize_final(client, full_text: str) -> dict:
    """One cheap text-only call: classify the document AND summarise it.

    Returns {"overview", "is_code", "language", "extension"}.
    """
    fallback = {"overview": "", "is_code": False, "language": "", "extension": ""}
    if not full_text.strip():
        return fallback
    msg = client.messages.create(
        model=MODEL,
        max_tokens=1024,
        system=FINALIZE_SYSTEM_PROMPT,
        messages=[{"role": "user", "content": full_text}],
    )
    raw = "".join(getattr(b, "text", "") for b in msg.content).strip()
    data = _parse_json(raw)
    if data is None:
        # Couldn't parse — keep the raw text as the overview, treat as non-code.
        return {**fallback, "overview": raw}
    return {
        "overview": str(data.get("overview", "")).strip(),
        "is_code": bool(data.get("is_code", False)),
        "language": str(data.get("language", "")).strip(),
        "extension": str(data.get("extension", "")).strip().lstrip(".").lower(),
    }


def _overlap_len(a: list, b: list, min_overlap: int = 2, max_check: int = 300) -> int:
    """Largest k such that the last k lines of a equal the first k lines of b
    (compared with trailing whitespace ignored). 0 if no run of >= min_overlap.
    Used to drop the duplicated region where two scrolled screenshots overlap."""
    limit = min(len(a), len(b), max_check)
    for k in range(limit, min_overlap - 1, -1):
        if [x.rstrip() for x in a[-k:]] == [x.rstrip() for x in b[:k]]:
            return k
    return 0


import re as _re

_FENCE_RE = _re.compile(r"```[^\n`]*\n(.*?)```", _re.S)
# leading line-number gutter: digits, an optional gutter glyph (Eclipse fold marker,
# middot, colon), then whitespace — e.g. "12  ", "5⊖ ", "3: "
_GUTTER_RE = _re.compile(r"^[ \t]*\d{1,4}[ \t\u00b7:.\u2296\u2299\u25cb\u2d54]?[ \t]+")


def _strip_gutter(text: str) -> str:
    lines = text.split("\n")
    nonempty = [l for l in lines if l.strip()]
    if nonempty and sum(1 for l in nonempty if _GUTTER_RE.match(l)) >= 0.6 * len(nonempty):
        lines = [_GUTTER_RE.sub("", l, count=1) if _GUTTER_RE.match(l) else l for l in lines]
    return "\n".join(lines)


def _indent_score(t: str) -> int:
    return sum(len(l) - len(l.lstrip()) for l in t.splitlines())


def _dedup_best(items: list) -> list:
    """Drop items that are the same code modulo whitespace; keep the best-indented
    copy of each, in first-seen order."""
    groups, order = {}, []
    for c in items:
        key = _re.sub(r"\s+", "", c)
        if not key:
            continue
        if key not in groups:
            groups[key] = c
            order.append(key)
        elif _indent_score(c) > _indent_score(groups[key]):
            groups[key] = c
    return [groups[k] for k in order]


_CODE_KW = _re.compile(
    r"^(public|private|protected|class|interface|abstract|import|package|def|function|"
    r"const|let|var|int|float|double|char|bool|void|static|final|return|new|struct|enum|"
    r"namespace|using|include|from|export|func|fn|type)\b")


def _strip_md_headers(text: str) -> str:
    """Drop Markdown ATX header lines a screen-reader adds over code — only when they're
    clearly code artifacts: a bare filename, a line starting with a code keyword, or text
    that appears in a real code line. Preserves ordinary comments and prose headers, and
    C preprocessor directives (which have no space after '#')."""
    lines = text.split("\n")
    codeset = [l.strip() for l in lines if l.strip() and not l.lstrip().startswith("#")]
    out = []
    for l in lines:
        m = _re.match(r"^\s*#{1,6}\s+(.+?)\s*$", l)
        if m:
            h = m.group(1).strip()
            is_file = bool(_re.match(r"^[\w./-]+\.[A-Za-z0-9]{1,5}$", h))
            is_kw = bool(_CODE_KW.match(h))
            in_code = len(h) > 4 and any(h in c or c in h for c in codeset)
            if is_file or is_kw or in_code:
                continue
        out.append(l)
    return "\n".join(out)


def clean_source(text: str) -> str:
    """Turn a raw OCR'd code extraction into compiler-ready source: unwrap markdown
    code fences (dropping ```lang, ``` and any # headers/prose outside them), remove
    an editor line-number gutter, and drop identical duplicate blocks. Faithful — it
    only strips transcription/formatting noise, never changes the code itself."""
    if not text or not text.strip():
        return text or ""
    blocks = _FENCE_RE.findall(text)
    if blocks:
        cleaned = [_strip_gutter(b).strip("\n") for b in blocks]
        text = "\n\n".join(_dedup_best(cleaned) or cleaned)
    else:
        text = _re.sub(r"^[ \t]*```.*$", "", text, flags=_re.M)  # stray fence lines
        text = _strip_gutter(text)
    text = _strip_md_headers(text)
    return text.strip("\n")


def merge_frames(raw_parts: list):
    """Clean each frame (fences/gutters/dup blocks), collapse frames that are the
    same code (keeping the best-indented copy), stitch overlapping ones. Returns
    (merged_code, clean_parts)."""
    cleaned = [clean_source(r) for r in raw_parts]
    parts = _dedup_best(cleaned) or [c for c in cleaned if c.strip()]
    return stitch_parts(parts), parts


def stitch_parts(parts: list) -> str:
    """Join per-image text, merging the overlap between consecutive chunks so
    scroll captures don't repeat their shared lines."""
    merged: list = []
    for part in parts:
        lines = part.split("\n")
        while lines and not lines[0].strip():
            lines.pop(0)
        while lines and not lines[-1].strip():
            lines.pop()
        if not lines:
            continue
        if not merged:
            merged = lines
            continue
        k = _overlap_len(merged, lines)
        merged += lines[k:] if k else [""] + lines
    return "\n".join(merged)


def analyse_incremental(client, image_paths: list, cache_dir: Path = None) -> dict:
    """Per-image extraction (+ content-hash cache) -> stitched text -> overview.

    Returns the same {"explanation", "extracted_text"} shape as analyse_images,
    so build_docx() and the callers work unchanged.
    """
    parts = []
    n = len(image_paths)
    for i, path in enumerate(image_paths, 1):
        data = path.read_bytes()
        digest = hashlib.sha256(data).hexdigest()
        cache_file = (cache_dir / f"{digest}.md") if cache_dir is not None else None

        if cache_file is not None and cache_file.exists():
            print(f"  [{i}/{n}] {path.name} (cached)")
            text = cache_file.read_text()
        else:
            print(f"  [{i}/{n}] reading {path.name}...")
            try:
                text = extract_one(client, path)
            except Exception as exc:  # noqa: BLE001 - skip a bad frame, keep the rest
                print(f"      (skipped — {type(exc).__name__}: {exc})", file=sys.stderr)
                continue
            if cache_file is not None:
                cache_dir.mkdir(parents=True, exist_ok=True)
                cache_file.write_text(text)

        if text.strip():
            parts.append(text.strip())

    full_text = stitch_parts(parts)
    if not full_text.strip():
        return {"explanation": "", "extracted_text": "",
                "is_code": False, "language": "", "extension": ""}

    print("  classifying + writing overview...")
    meta = synthesize_final(client, full_text)
    return {
        "explanation": meta["overview"],
        "extracted_text": full_text,
        "is_code": meta["is_code"],
        "language": meta["language"],
        "extension": meta["extension"],
    }




def cache_path_for(path: Path, cache_dir: Path) -> Path:
    """Content-addressed cache location for an image's extracted text."""
    digest = hashlib.sha256(path.read_bytes()).hexdigest()
    return cache_dir / f"{digest}.md"


def extract_to_cache(client, path: Path, cache_dir: Path) -> None:
    """Extract one image's text and cache it (no-op if already cached).

    Used for background pre-extraction at capture time; analyse_incremental()
    later finds the cache hit and skips the API call. Same hash/key scheme as
    analyse_incremental so the two share one cache.
    """
    cf = cache_path_for(path, cache_dir)
    if cf.exists():
        return
    text = extract_one(client, path)
    cache_dir.mkdir(parents=True, exist_ok=True)
    cf.write_text(text)


# ── Code fix loop (Milestone 7, Layer C) ─────────────────────────────────────────

FIX_SYSTEM_PROMPT = (
    "You are correcting a source file that was transcribed from screenshots and "
    "failed to compile. You are given the language, the compiler/parser errors, "
    "and the current code. The errors are TRANSCRIPTION mistakes (a mis-read "
    "character, a missing bracket/colon/semicolon, a wrong quote, a broken indent).\n"
    "Fix ONLY what the errors point to, with the MINIMAL change needed. Do NOT add, "
    "remove, rename, or invent functionality, imports, comments, or logic that is "
    "not clearly required to resolve the specific reported error. Preserve the "
    "original code exactly everywhere else. If a fix is genuinely ambiguous, leave "
    "that line unchanged rather than guessing.\n"
    "Return ONLY the corrected, complete source file — no commentary, no markdown."
)


def fix_source(client, code: str, language: str, errors: str) -> str:
    """One API call: return a corrected version of the code given compiler errors."""
    msg = client.messages.create(
        model=MODEL,
        max_tokens=4096,
        system=FIX_SYSTEM_PROMPT,
        messages=[{"role": "user", "content": (
            f"Language: {language or 'unknown'}\n\n"
            f"Compiler/parser errors:\n{errors}\n\n"
            f"Current code:\n{code}"
        )}],
    )
    return "".join(getattr(b, "text", "") for b in msg.content).strip()


# ── Document builder ───────────────────────────────────────────────────────────

# Converts Markdown returned by Claude into native Word paragraph styles.
def markdown_to_docx(doc, md_text: str) -> None:
    if not md_text.strip():
        doc.add_paragraph("(No text content detected.)")
        return

    for line in md_text.splitlines():
        if line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=3)
        elif line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=2)
        elif line.startswith("# "):
            doc.add_heading(line[2:].strip(), level=1)
        elif line.startswith("- ") or line.startswith("* "):
            doc.add_paragraph(line[2:].strip(), style="List Bullet")
        elif len(line) > 2 and line[0].isdigit() and line[1:3] in (". ", ") "):
            doc.add_paragraph(line[3:].strip(), style="List Number")
        elif line.strip() == "":
            doc.add_paragraph("")
        else:
            doc.add_paragraph(line.strip())


# Writes only the extracted text to a Word doc — no headers or labels.
def build_docx(result: dict) -> "Document":
    from docx import Document

    doc = Document()
    extracted = result.get("extracted_text", "").strip()
    if extracted:
        markdown_to_docx(doc, extracted)
    else:
        doc.add_paragraph("(No text content was extracted from the images.)")
    return doc
