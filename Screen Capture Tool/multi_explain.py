"""Milestone 2.5 — Multi-image explain + optional doc export.

How it works end-to-end:
  1. The user passes one or more image files as arguments.
  2. All images are base64-encoded and sent to Claude in a SINGLE API call,
     so Claude sees them as one continuous document — not isolated screenshots.
  3. While Claude processes, a spinner with rotating status messages keeps
     the user informed (streaming runs hidden in the background).
  4. Once done, the terminal prints:
       - OVERVIEW: a plain-English description of what is across all the images.
       - EXTRACTED TEXT: all visible text stitched together in document order.
  5. The user is asked whether to save the extracted text to a .docx file.
     If yes, the file is saved into the same folder the images came from.

Prerequisites:
  pip install -r requirements.txt
  ANTHROPIC_API_KEY=sk-ant-... in .env (or exported in shell)

Run it:
  python multi_explain.py ~/Desktop/SSTest/*.png
  python multi_explain.py img1.png img2.png img3.jpg
"""

import argparse
import base64
import itertools
import json
import os
import sys
import threading
import time
from pathlib import Path

# The vision-capable model used for all API calls.
MODEL = "claude-sonnet-4-6"

# System prompt — tells Claude to treat all images as one continuous document
# and return a strict JSON object with two keys. Strict JSON makes it easy to
# split the explanation from the extracted text without fragile text parsing.
SYSTEM_PROMPT = (
    "You are a screen-reading assistant that analyses a sequence of screenshots as one continuous piece of content.\n"
    "The images are ordered and together represent a single document, page, or screen flow.\n"
    "Return a JSON object with exactly two keys:\n"
    "\n"
    '  "explanation": A detailed plain-English overview of what is shown across all the images combined. '
    "Lead with the content type (e.g. 'Spreadsheet:', 'Document:', 'Code editor:'), then describe "
    "the full picture — layout, key elements, how the images relate to each other.\n"
    "\n"
    '  "extracted_text": All visible text from all images stitched together in order as one '
    "continuous document, preserving the original structure throughout. Use Markdown:\n"
    "    - Headings → # / ## / ### etc.\n"
    "    - Bullet lists → -\n"
    "    - Numbered lists → 1. 2. 3.\n"
    "    - Tables → Markdown pipe tables\n"
    "    - Plain paragraphs → plain paragraphs separated by blank lines\n"
    "    Continue numbering, headings, and structure naturally across images as if they are "
    "one document. Do not restart or add separators between images.\n"
    '    If there is no meaningful text, use an empty string "".\n'
    "\n"
    "Return ONLY the raw JSON object. No code fences, no extra keys, no commentary."
)

# The user-facing instruction sent alongside the images.
USER_PROMPT = "Analyse all these screenshots as one continuous document and return the JSON as instructed."


# ── Environment ────────────────────────────────────────────────────────────────

# Loads ANTHROPIC_API_KEY from a .env file into os.environ so the rest of the
# script can read it. If the key is already exported in the shell this is a
# no-op. If python-dotenv is not installed we skip silently — no crash.
def load_env() -> None:
    try:
        from dotenv import load_dotenv
        load_dotenv()
    except ImportError:
        pass


# ── Spinner ────────────────────────────────────────────────────────────────────

# Status messages that rotate in the terminal while Claude is working.
# They describe what is happening at a high level and keep the user engaged
# during what can be a 15-30 second wait for a multi-image analysis.
PROCESSING_MESSAGES = [
    "Reading pixel data across all images...",
    "Identifying content types and layout structures...",
    "Cross-referencing text regions between images...",
    "Analysing visual hierarchy and document flow...",
    "Stitching content together into a single document...",
    "Extracting and preserving text formatting...",
    "Resolving structure across image boundaries...",
    "Almost there — finalising the analysis...",
]


# Runs on a background thread while the API call is in progress.
# Prints an animated braille spinner + a rotating status message so the user
# knows the tool is actively working. Stops as soon as stop_event is set by
# the main thread, then clears its line so the output stays clean.
def _spinner(stop_event: threading.Event) -> None:
    spinner = itertools.cycle(["⠋", "⠙", "⠹", "⠸", "⠼", "⠴", "⠦", "⠧", "⠇", "⠏"])
    messages = itertools.cycle(PROCESSING_MESSAGES)
    current_msg = next(messages)
    msg_timer = time.time()

    while not stop_event.is_set():
        frame = next(spinner)
        # \r returns to the start of the line so the spinner animates in place.
        print(f"\r  {frame}  {current_msg}   ", end="", flush=True)
        time.sleep(0.1)

        # Advance to the next status message every 3 seconds.
        if time.time() - msg_timer > 3:
            current_msg = next(messages)
            msg_timer = time.time()

    # Wipe the spinner line so it does not bleed into the output below.
    print("\r" + " " * 70 + "\r", end="", flush=True)


# ── API call ───────────────────────────────────────────────────────────────────

# Core API function. Encodes every image to base64 and sends them ALL in one
# Messages API call so Claude has full cross-image context. Sending separately
# would lose the thread between images and produce disconnected results.
# Streams the response silently while the spinner runs, then parses the JSON.
# Returns a dict: {"explanation": str, "extracted_text": str}.
def analyse_images(client, image_paths: list) -> dict:
    # Build the content list: one image block per file, instruction text last.
    content = []
    for path in image_paths:
        # base64-encode the raw file bytes — required by the Messages API
        # for inline image uploads.
        b64 = base64.standard_b64encode(path.read_bytes()).decode()
        content.append({
            "type": "image",
            "source": {
                "type": "base64",
                "media_type": _media_type(path),  # e.g. "image/png"
                "data": b64,
            },
        })
    # The text instruction must come after all the images.
    content.append({"type": "text", "text": USER_PROMPT})

    # Launch the spinner on a daemon thread so it runs alongside the blocking
    # API call. daemon=True ensures it cannot prevent the process from exiting.
    stop_event = threading.Event()
    spinner_thread = threading.Thread(target=_spinner, args=(stop_event,), daemon=True)
    spinner_thread.start()

    # Stream the response and accumulate it in `raw`. We hide the stream from
    # the terminal because it is raw JSON — unreadable to the user. The spinner
    # is what keeps them informed during this wait.
    raw = ""
    try:
        with client.messages.stream(
            model=MODEL,
            max_tokens=4096,
            system=SYSTEM_PROMPT,
            messages=[{"role": "user", "content": content}],
        ) as stream:
            for text in stream.text_stream:
                raw += text  # silently collect each chunk
    finally:
        # Always stop the spinner — even if the API call raises an exception.
        stop_event.set()
        spinner_thread.join()

    raw = raw.strip()

    # Claude sometimes wraps the JSON in ```json ... ``` code fences despite
    # being told not to. Strip them so json.loads always receives clean input.
    if raw.startswith("```"):
        raw = raw.split("\n", 1)[-1]           # remove the opening fence line
        raw = raw.rsplit("```", 1)[0].strip()  # remove the closing fence

    # First attempt: parse the full response as JSON directly.
    try:
        return json.loads(raw)
    except json.JSONDecodeError:
        pass

    # Second attempt: find the outermost { ... } block and parse just that.
    # This handles cases where Claude adds a preamble sentence or trailing note
    # around the JSON object despite being told not to.
    start = raw.find("{")
    end = raw.rfind("}")
    if start != -1 and end != -1 and end > start:
        try:
            return json.loads(raw[start:end + 1])
        except json.JSONDecodeError:
            pass

    # Final fallback: parsing failed entirely. Warn the user so they know
    # something went wrong, and show the raw response as the explanation.
    print("\n[Warning] Could not parse a structured response — showing raw output.", file=sys.stderr)
    return {"explanation": raw, "extracted_text": ""}


# Helper that maps a file extension to the MIME type the Messages API expects
# for inline base64 images. Defaults to image/png for unknown extensions.
def _media_type(path: Path) -> str:
    return {
        ".jpg": "image/jpeg",
        ".jpeg": "image/jpeg",
        ".png": "image/png",
        ".gif": "image/gif",
        ".webp": "image/webp",
    }.get(path.suffix.lower(), "image/png")


# ── Document builder ───────────────────────────────────────────────────────────

# Converts Markdown-formatted text (as returned by Claude) into native Word
# paragraph styles inside a python-docx Document. We parse line by line and map
# each Markdown pattern to the matching docx style so the document looks
# properly structured rather than displaying raw Markdown syntax.
def markdown_to_docx(doc, md_text: str) -> None:
    if not md_text.strip():
        doc.add_paragraph("(No text content detected.)")
        return

    for line in md_text.splitlines():
        if line.startswith("### "):
            # Level 3 heading — e.g. ### Sub-section
            doc.add_heading(line[4:].strip(), level=3)
        elif line.startswith("## "):
            # Level 2 heading — e.g. ## Section
            doc.add_heading(line[3:].strip(), level=2)
        elif line.startswith("# "):
            # Level 1 heading — e.g. # Title
            doc.add_heading(line[2:].strip(), level=1)
        elif line.startswith("- ") or line.startswith("* "):
            # Unordered bullet point
            doc.add_paragraph(line[2:].strip(), style="List Bullet")
        elif len(line) > 2 and line[0].isdigit() and line[1:3] in (". ", ") "):
            # Ordered list item — handles both "1. " and "1) " formats
            doc.add_paragraph(line[3:].strip(), style="List Number")
        elif line.startswith("|"):
            # Markdown table row — render cells joined by a readable separator
            cells = [c.strip() for c in line.strip("|").split("|")]
            # Skip pure separator rows like |---|---| which carry no content
            if all(set(c) <= {"-", ":"} for c in cells):
                continue
            doc.add_paragraph("  |  ".join(cells))
        elif line.strip() == "":
            # Blank line — preserves paragraph spacing in the document
            doc.add_paragraph("")
        else:
            # Plain paragraph text
            doc.add_paragraph(line.strip())


# Assembles the final Word document from the result dict returned by
# analyse_images. Only the extracted_text field is written — no overviews,
# image labels, or headers. The goal is a clean document that reads as if
# typed directly from the original source material.
def build_docx(result: dict) -> "Document":
    from docx import Document

    doc = Document()

    extracted = result.get("extracted_text", "").strip()
    if extracted:
        markdown_to_docx(doc, extracted)
    else:
        doc.add_paragraph("(No text content was extracted from the images.)")

    return doc


# ── Entry point ────────────────────────────────────────────────────────────────

# Main orchestrator — runs the full pipeline from CLI arguments through to
# terminal output and optional .docx export. Returns 0 on success, 1 on error.
def main() -> int:

    # Step 1: Parse CLI arguments.
    # nargs="+" requires at least one image path; passing none prints usage.
    parser = argparse.ArgumentParser(
        description="Explain screenshots in the terminal, optionally export extracted text to .docx."
    )
    parser.add_argument("images", nargs="+", metavar="IMAGE", help="Image file paths to process.")
    args = parser.parse_args()

    # Step 2: Load the API key from .env (or fall through if already in shell).
    load_env()

    api_key = os.environ.get("ANTHROPIC_API_KEY")
    if not api_key:
        # Exit early — without a key the API call will always fail.
        print("ANTHROPIC_API_KEY is not set. Add it to .env.", file=sys.stderr)
        return 1

    # Step 3: Import the Anthropic SDK and create the client.
    # Imported here (not at the top) so a missing package gives a friendly
    # message instead of a raw ImportError on startup.
    try:
        import anthropic
    except ImportError:
        print("anthropic not installed. Run: pip install -r requirements.txt", file=sys.stderr)
        return 1

    # The client holds the API key and manages the HTTP connection.
    client = anthropic.Anthropic(api_key=api_key)

    # Step 4: Validate all image paths before making any API calls.
    # Fail fast here rather than halfway through a batch.
    paths = [Path(p) for p in args.images]
    for p in paths:
        if not p.exists():
            print(f"File not found: {p}", file=sys.stderr)
            return 1

    # The output .docx is always saved into the same folder as the source images.
    # Using the first image's parent keeps output co-located with the input.
    output_folder = paths[0].parent

    # Step 5: Send all images to Claude in one API call.
    # The spinner animates while we wait; the response is collected silently.
    print(f"\nSending {len(paths)} image(s) to Claude...\n")
    try:
        result = analyse_images(client, paths)
    except Exception as exc:
        print(f"FAILED — {type(exc).__name__}: {exc}", file=sys.stderr)
        return 1

    # Step 6: Print the results to the terminal in plain English.
    # The overview describes what Claude saw across all images combined.
    print(f"\n{'=' * 60}")
    print("OVERVIEW")
    print("=" * 60)
    print(result.get("explanation", "").strip())

    # The extracted text is shown as-is — Markdown is readable in a terminal.
    extracted = result.get("extracted_text", "").strip()
    if extracted:
        print(f"\n{'-' * 60}")
        print("EXTRACTED TEXT")
        print("-" * 60)
        print(extracted)
    else:
        print("\nNo text content detected in these images.")

    # Step 7: Ask whether to save a Word document.
    # The doc is only created on demand — no point writing a file the user
    # does not need. If yes, the extracted text is written in proper Word
    # formatting and saved next to the source images as report.docx.
    print(f"\n{'=' * 60}")
    answer = input("Save extracted text to a Word document? [y/n]: ").strip().lower()

    if answer == "y":
        # Check python-docx is available before attempting to build the doc.
        try:
            from docx import Document  # noqa: F401
        except ImportError:
            print("python-docx not installed. Run: pip install -r requirements.txt", file=sys.stderr)
            return 1

        # Guard against saving an empty document when nothing was extracted.
        if not extracted:
            print("Nothing to save — no text was extracted.")
            return 0

        # Build and save the document. report.docx overwrites any previous run.
        doc = build_docx(result)
        out_path = output_folder / "report.docx"
        doc.save(out_path)
        print(f"Saved: {out_path.resolve()}")
    else:
        print("No document saved.")

    return 0


if __name__ == "__main__":
    raise SystemExit(main())
